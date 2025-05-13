from flask import Flask, request, send_file, jsonify, render_template
from docx import Document
from docx.shared import Inches
import pandas as pd
import tempfile
import base64
import os
from zipfile import ZipFile

app = Flask(__name__)
# ✅ 1. Bu yerga qo‘shing:
def insert_band_paragraphs(doc, bandlar):
    for i, par in enumerate(doc.paragraphs):
        if '{{barcha_bandlar}}' in par.text:
            parent = par._element.getparent()
            idx = parent.index(par._element)
            parent.remove(par._element)
            for band in bandlar:
                new_par = doc.add_paragraph(f"- {band['matn']}")
                parent.insert(idx, new_par._element)
                idx += 1
            print("✅ Bandlar hujjatga qo‘shildi")
            break
# 1. Bitta paragraph uchun almashtirish
def replace_placeholders(paragraph, replacements):
    full_text = ''.join(run.text for run in paragraph.runs)
    if any(key in full_text for key in replacements):
        for key, val in replacements.items():
            full_text = full_text.replace(key, val)
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full_text

# 2. Butun hujjat bo'ylab barcha joylarda placeholderlarni almashtiruvchi funksiyani qo‘shamiz:
def replace_all_placeholders(doc, replacements):
    # Oddiy paragraflar
    for paragraph in doc.paragraphs:
        replace_placeholders(paragraph, replacements)

    # Jadval ichidagi paragraflar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders(paragraph, replacements)

    # Header va footer paragraflari (sarlavha/poytaxt qismlarda bo'lishi mumkin)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_placeholders(paragraph, replacements)

def process_docx(template_path, replacements, bandlar=None, include_band_table=False, imzo_map=None):
    doc = Document(template_path)

    replace_all_placeholders(doc, replacements)  # <<< Faqat mana shu qator orqali har joyni qamrab oladi ✅

    if bandlar and any("{{barcha_bandlar}}" in p.text for p in doc.paragraphs):
        insert_band_paragraphs(doc, bandlar)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    replace_placeholders(par, replacements)

    if include_band_table and bandlar and doc.tables:
        band_table = doc.tables[0]
        while len(band_table.rows) > 1:
            band_table._tbl.remove(band_table.rows[1]._tr)
        for i, band in enumerate(bandlar, start=1):
            row = band_table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = band.get("matn", "")
            row.cells[2].text = band.get("muddat", "")

    if imzo_map:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        for key, data_url in imzo_map.items():
                            if key in par.text and data_url:
                                par.text = ""
                                try:
                                    img_bytes = base64.b64decode(data_url.split(",")[1])
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_file:
                                        img_file.write(img_bytes)
                                        img_path = img_file.name
                                    par.add_run().add_picture(img_path, width=Inches(1.5))
                                except Exception as e:
                                    print(f"Xatolik: {key} - {e}")

    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_path)
    return temp_path

@app.route("/")
def index():
    return render_template("form.html")

@app.route("/generate", methods=["POST"])
def generate():
    data = request.get_json()

    # Ma'lumotlarni ajratib olish
    ism = data.get("ism")
    lavozim = data.get("lavozim")
    qoidabuzarlik = data.get("qoidabuzarlik")
    bandlar = data.get("bandlar", [])
    imzo_inspektor = data.get("imzo_inspektor")
    imzo_obyekt = data.get("imzo_obyekt")
    imzo_qatnashganlar = data.get("imzo_qatnashganlar")
    obyekt_nomi = data.get("obyekt_nomi", "")
    obyekt_rahbari = data.get("obyekt_rahbari", "")
    DYON_organi = data.get("DYON_organi", "")
    tekshiruvchilar = data.get("tekshiruvda_qatnashganlar", "")
    sanasi = data.get("sanasi", "")
    nazorat_sana = data.get("nazorat_sana", "")
    termiz = data.get("termiz", "")

    # === BANDLARni matn ko‘rinishiga aylantirish funksiyasi ===
    def bandlar_to_text(bandlar):
        if not bandlar:
            return "Bandlar kiritilmagan."
        return "\n".join([
            f"{i+1}) {b['matn']} (Bajarish muddati: {b['muddat']})"
            for i, b in enumerate(bandlar)
        ])

    replacements = {
        "{{ism}}": ism,
        "{{lavozim}}": lavozim,
        "{{qoidabuzarlik}}": qoidabuzarlik,
        "{{Obyekt_nomi}}": obyekt_nomi,
        "{{Obyekt_rahbari}}": obyekt_rahbari,
        "{{DYON_organi}}": DYON_organi,
        "{{Tekshiruvda_qatnashganlar}}": tekshiruvchilar,
        "{{Inspektor_lavozimi_Ismi}}": lavozim + " " + ism,
        "{{kun.oy.yil}}": sanasi,
        "{{yil.oy.yil}}": nazorat_sana,
        "{{Termiz}}": termiz,
    }
    
    # === Hujjatlarni yuklash ===
    doc1 = Document("hujjat1.docx")
    doc2 = Document("hujjat2.docx")
   
    imzo_map1 = {
        "{{IMZO_INSPEKTOR}}": imzo_inspektor,
        "{{IMZO_OBYEKT}}": imzo_obyekt
    }
    imzo_map2 = {
        "{{IMZO_INSPEKTOR}}": imzo_inspektor,
        "{{IMZO_OBYEKT}}": imzo_obyekt,
        "{{IMZO_QATNASHGANLAR}}": imzo_qatnashganlar
    }

    doc1 = process_docx("hujjat1.docx", replacements, bandlar, include_band_table=True, imzo_map=imzo_map1)
    doc2 = process_docx("hujjat2.docx", replacements, bandlar, include_band_table=False, imzo_map=imzo_map2)

    zip_path = tempfile.NamedTemporaryFile(delete=False, suffix=".zip").name
    with ZipFile(zip_path, "w") as zipf:
        zipf.write(doc1, arcname="Yozma_Korsatma.docx")
        zipf.write(doc2, arcname="Dalolatnoma.docx")

    return send_file(zip_path, as_attachment=True, download_name="tayyor_hujjatlar.zip")

@app.route("/search-band", methods=["POST"])
def search_band():
    query = request.json.get("query", "").strip()
    if not query:
        return jsonify({"natija": []})

    try:
        df = pd.read_excel("bandlar.xlsx")
        df = df.dropna()
        topilgan = []
        for _, row in df.iterrows():
            matn = str(row.get("Hujjat bandi", "")).strip()
            if query.lower() in matn.lower():
                topilgan.append({"matn": matn, "muddat": ""})
        return jsonify({"natija": topilgan[:5]})
    except Exception as e:
        return jsonify({"xatolik": str(e), "natija": []})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    app.run(host='0.0.0.0', port=port)

